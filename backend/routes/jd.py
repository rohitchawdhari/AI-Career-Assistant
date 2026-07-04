import os
import io
import re
import json
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document

from rag.vector_store import vector_store
from resume.skill_extractor import extract_skills
from resume.jd_matcher import analyze_jd_match
from resume.parser import parse_file
from services.gemini_service import model

router = APIRouter()


class JDRequest(BaseModel):
    job_description: str


class JDAIRequest(BaseModel):
    job_description: str


# Keep existing endpoint for backward compatibility
@router.post("/analyze-jd")
async def analyze_jd(request: JDRequest):
    resume_text = "\n".join(vector_store.chunks)
    resume_skills = extract_skills(resume_text)
    result = analyze_jd_match(resume_skills, request.job_description)
    return result


# Keep existing endpoint for backward compatibility
@router.post("/analyze-jd-ai")
async def analyze_jd_ai(request: JDAIRequest):
    if not vector_store.chunks:
        return {"error": "Please upload a resume first."}

    resume_text = vector_store.get_full_text()
    jd_text = request.job_description

    prompt = f"""
    You are an expert recruiter and career consultant.
    Compare the following candidate resume text with the job description (JD) text.
    Provide a structured JSON response with:
    1. "match_score": A realistic match score from 0 to 100 based on core skills, experience level alignment, and role fit.
    2. "matched_skills": A list of skills present in both the resume and the JD.
    3. "missing_skills": A list of important skills mentioned in the JD but missing or weak in the resume.
    4. "keyword_gaps": A list of specific industry keywords/buzzwords in the JD that are not in the resume.
    5. "suggestions": A list of actionable resume improvements to increase alignment with this specific JD.

    Resume Text:
    \"\"\"{resume_text}\"\"\"

    Job Description:
    \"\"\"{jd_text}\"\"\"

    Your output must be a valid JSON object. Do not wrap it in markdown code blocks like ```json ... ```. Just return the JSON object directly.
    Example format:
    {{
        "match_score": 82,
        "matched_skills": ["Python", "SQL", "Git"],
        "missing_skills": ["AWS", "Docker"],
        "keyword_gaps": ["Cloud Infrastructure", "Containerization"],
        "suggestions": ["Highlight DevOps achievements in your work history", "Mention any AWS projects or certifications you have"]
    }}
    """
    try:
        response = model.generate_content(prompt)
        clean_text = response.text.strip()
        if clean_text.startswith("```json"):
            clean_text = clean_text[7:]
        if clean_text.endswith("```"):
            clean_text = clean_text[:-3]
        clean_text = clean_text.strip()

        analysis = json.loads(clean_text)
        return analysis
    except Exception as e:
        print(f"Error in Gemini JD matching: {e}")
        fallback = analyze_jd_match(extract_skills(resume_text), jd_text)
        return {
            "match_score": fallback["match_score"],
            "matched_skills": fallback["matched_skills"],
            "missing_skills": fallback["missing_skills"],
            "keyword_gaps": ["AI comparison failed, falling back to keyword comparison"],
            "suggestions": ["Ensure your technical skill section contains matches from the JD."]
        }


# Keep existing endpoint for backward compatibility
@router.post("/upload-jd")
async def upload_jd(file: UploadFile = File(...)):
    os.makedirs("uploads", exist_ok=True)
    file_path = os.path.join("uploads", file.filename)

    with open(file_path, "wb") as f:
        f.write(await file.read())

    try:
        jd_text = parse_file(file_path)
    except Exception as e:
        return {"error": f"Failed to parse JD file: {str(e)}"}

    if not jd_text.strip():
        return {"error": "Extracted job description text is empty."}

    resume_text = "\n".join(vector_store.chunks)
    resume_skills = extract_skills(resume_text)
    result = analyze_jd_match(resume_skills, jd_text)
    result["extracted_text"] = jd_text
    return result


# ========================================================
# NEW ENDPOINTS FOR COMPREHENSIVE JD MATCH & OPTIMIZATION
# ========================================================

@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    """Extracts text from PDF, DOC, or DOCX files for Job Descriptions or Resume uploads."""
    os.makedirs("uploads", exist_ok=True)
    file_path = os.path.join("uploads", file.filename)

    with open(file_path, "wb") as f:
        f.write(await file.read())

    try:
        extracted_text = parse_file(file_path)
        if not extracted_text or not extracted_text.strip():
            raise HTTPException(status_code=400, detail="The uploaded file appears to contain no text.")
        return {
            "text": extracted_text,
            "filename": file.filename
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {str(e)}")


@router.post("/analyze-resume-jd")
async def analyze_resume_jd(request: JDRequest):
    """Performs a deep comparison of the uploaded resume and job description using detailed metrics."""
    if not vector_store.chunks:
        raise HTTPException(status_code=400, detail="Please upload a resume first.")

    resume_text = vector_store.get_full_text()
    jd_text = request.job_description

    prompt = f"""
    You are an expert ATS (Applicant Tracking System) parser and technical recruiter.
    Analyze the candidate's resume and compare it systematically against the target Job Description (JD).
    
    You must compute a highly realistic, non-random evaluation of how well the candidate fits the requirements.
    
    Evaluate the following metrics:
    - Keyword relevance: Presence of target terms.
    - Skill matching: Technical or soft skills overlap.
    - Experience relevance: Work history fit.
    - Education match: Degree & certification alignment.
    - Project match: Relevant builds or initiatives.
    - Formatting: Structure, spacing, sections, readability.
    - Action verbs: Impactful wording vs passive phrases.
    - Achievements: Quantifiable metrics.
    
    Resume Text:
    \"\"\"{resume_text}\"\"\"

    Job Description:
    \"\"\"{jd_text}\"\"\"

    Return a raw JSON object only (do NOT wrap it in markdown code blocks like ```json ... ```).
    The response must contain exactly these keys:
    1. "ats_match_score": A realistic match score from 0 to 100 based on the overlaps.
    2. "resume_score": Standalone resume score from 0 to 100 reflecting its overall formatting and strength.
    3. "keyword_match": A list of matching key phrases or tools found in both.
    4. "missing_keywords": A list of target keywords from the JD that are not in the resume.
    5. "matching_skills": A list of matching technical or soft skills.
    6. "missing_skills": A list of missing skills mentioned in the JD.
    7. "experience_match": A detailed 1-2 sentence evaluation of the work history alignment.
    8. "education_match": A 1-2 sentence evaluation of how well the degrees match the JD requirements.
    9. "project_match": A 1-2 sentence evaluation of the projects alignment.
    10. "certification_match": A 1-2 sentence evaluation of the certifications/credentials fit.
    11. "overall_compatibility_score": An overall compatibility score from 0 to 100.
    12. "suggestions": A nested JSON object containing:
        - "missing_skills": A list of skills to learn or list.
        - "missing_technologies": A list of technologies/tools to study.
        - "weak_summary": A list of suggestions to improve the resume summary (empty list if strong).
        - "weak_experience": A list of bullet point improvements for work experience.
        - "weak_projects": A list of recommendations to optimize projects.
        - "weak_achievements": A list of tips to add metrics or quantifiable achievements.
        - "weak_keywords": A list of keywords to include.
        - "missing_certifications": A list of relevant certifications to pursue.
        - "action_verbs": A list of strong action verbs to replace passive ones.
        - "grammar_improvements": A list of grammar or voice edits.
        - "formatting_suggestions": A list of formatting or styling changes.
    """

    try:
        response = model.generate_content(prompt)
        clean_text = response.text.strip()
        if clean_text.startswith("```json"):
            clean_text = clean_text[7:]
        if clean_text.endswith("```"):
            clean_text = clean_text[:-3]
        clean_text = clean_text.strip()

        analysis = json.loads(clean_text)
        return analysis
    except Exception as e:
        print(f"Error in analyze_resume_jd: {e}")
        # Fallback to deterministic model
        fallback = analyze_jd_match(extract_skills(resume_text), jd_text)
        return {
            "ats_match_score": fallback["match_score"],
            "resume_score": 70,
            "keyword_match": fallback["matched_skills"],
            "missing_keywords": fallback["missing_skills"][:5],
            "matching_skills": fallback["matched_skills"],
            "missing_skills": fallback["missing_skills"],
            "experience_match": "AI analysis fell back to keyword matching. Work experience overlap was evaluated on technical skills.",
            "education_match": "Educational requirements check was bypassed due to backend fallback.",
            "project_match": "Project alignment evaluation is pending a fresh AI connection.",
            "certification_match": "Certification matching was bypassed during fallback.",
            "overall_compatibility_score": fallback["match_score"],
            "suggestions": {
                "missing_skills": fallback["missing_skills"],
                "missing_technologies": fallback["missing_skills"],
                "weak_summary": ["Provide a concise summary tailored to the job description."],
                "weak_experience": ["Quantify achievements in your professional history using numbers and percentages."],
                "weak_projects": ["Include technologies and system metrics in project descriptions."],
                "weak_achievements": ["Add metrics showing the business value of your work."],
                "weak_keywords": fallback["missing_skills"][:3],
                "missing_certifications": [],
                "action_verbs": ["Led", "Developed", "Optimized", "Designed"],
                "grammar_improvements": [],
                "formatting_suggestions": ["Keep section headings uniform and ensure clean alignments."]
            }
        }


@router.post("/optimize-resume")
async def optimize_resume(request: JDRequest):
    """Rewrites the active resume to match the job description using advanced AI wording and keywords without changing facts."""
    if not vector_store.chunks:
        raise HTTPException(status_code=400, detail="Please upload a resume first.")

    resume_text = vector_store.get_full_text()
    jd_text = request.job_description

    prompt = f"""
    You are an expert resume writer and recruiter.
    Rewrite the candidate's resume to align it with the target Job Description (JD).
    
    CRITICAL RULES:
    1. Do NOT change factual information (do not invent experiences, jobs, companies, projects, or certifications).
    2. Only improve wording, structure, formatting, and keyword placement.
    3. Use the STAR method (Situation, Task, Action, Result) with strong action verbs.
    4. Highlight and integrate relevant keywords from the JD naturally.

    Resume Text:
    \"\"\"{resume_text}\"\"\"

    Job Description:
    \"\"\"{jd_text}\"\"\"

    Return a raw JSON object only (do NOT wrap it in markdown code blocks like ```json ... ```).
    The response must contain exactly these keys:
    1. "original_resume": The original resume content.
    2. "optimized_resume_markdown": The fully rewritten, optimized resume in clean, beautifully styled Markdown.
    3. "optimized_summary": The rewritten Professional Summary.
    4. "optimized_skills": The categorized technical skills list.
    5. "optimized_projects": The rewritten Projects section.
    6. "optimized_experience": The rewritten Work Experience bullet points.
    7. "added_keywords": A list of keywords from the JD that were successfully added or emphasized.
    8. "improved_bullets": A list of JSON objects, each with:
        - "before": The original bullet point text.
        - "after": The rewritten optimized bullet point text.
    9. "improved_summary_explanation": A brief description of how the profile summary was optimized.
    10. "ats_improvement": A description of the improvements made to increase ATS parsing alignment.
    11. "estimated_ats_increase": An estimated percentage points increase in the ATS score (integer, e.g. 15).
    """

    try:
        response = model.generate_content(prompt)
        clean_text = response.text.strip()
        if clean_text.startswith("```json"):
            clean_text = clean_text[7:]
        if clean_text.endswith("```"):
            clean_text = clean_text[:-3]
        clean_text = clean_text.strip()

        optimized_data = json.loads(clean_text)
        return optimized_data
    except Exception as e:
        print(f"Error in optimize_resume: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to optimize resume: {str(e)}")


class DownloadDOCXRequest(BaseModel):
    markdown: str


@router.post("/download-docx")
async def download_docx(request: DownloadDOCXRequest):
    """Generates a downloadable Microsoft Word document (.docx) from Markdown content."""
    markdown_content = request.markdown

    doc = Document()
    
    # Set standard styles
    title = doc.add_heading("AI Optimized Resume", 0)
    title.alignment = 1  # Centered

    lines = markdown_content.split("\n")
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue

        # Headings
        if line_strip.startswith("# "):
            doc.add_heading(line_strip[2:], level=1)
        elif line_strip.startswith("## "):
            doc.add_heading(line_strip[3:], level=2)
        elif line_strip.startswith("### "):
            doc.add_heading(line_strip[4:], level=3)
        # Bullet points
        elif line_strip.startswith("* ") or line_strip.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            text = line_strip[2:]
            
            # Simple bold parser (**bold**)
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        # Normal text
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", line_strip)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Save to a memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=optimized_resume.docx"}
    )