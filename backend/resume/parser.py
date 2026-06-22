import os
import re
from pypdf import PdfReader
from docx import Document

def parse_pdf(file_path):
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error parsing PDF: {e}")
    return text

def parse_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
    return text

def parse_doc(file_path):
    # First attempt: Using win32com.client if MS Word is installed (since OS is Windows)
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_path = os.path.abspath(file_path)
        doc = word.Documents.Open(abs_path)
        text = doc.Content.Text
        doc.Close()
        word.Quit()
        return text
    except Exception as e:
        print(f"win32com parsing failed or MS Word not installed, using fallback scanner: {e}")
        # Second attempt: Heuristic binary text extraction for OLE files
        try:
            with open(file_path, "rb") as f:
                data = f.read()
            
            # Legacy .doc files usually store UTF-16 or ASCII characters.
            # We find printable character chains.
            text_chunks = []
            
            # 1. UTF-16 strings (usually encoded as letter + null byte)
            utf16_matches = re.findall(b"(?:[\x20-\x7E\x0A\x0D]\x00){4,}", data)
            for m in utf16_matches:
                try:
                    text_chunks.append(m.decode("utf-16le", errors="ignore"))
                except Exception:
                    pass
                    
            # 2. ASCII strings
            ascii_matches = re.findall(b"[\x20-\x7E\x0A\x0D]{4,}", data)
            for m in ascii_matches:
                try:
                    text_chunks.append(m.decode("ascii", errors="ignore"))
                except Exception:
                    pass
            
            # Combine and clean up
            raw_text = "\n".join(text_chunks)
            lines = []
            for line in raw_text.splitlines():
                line = line.strip()
                if len(line) < 3:
                    continue
                # Simple cleanup: ensure lines have realistic word boundaries (mostly words, numbers, punctuation)
                alnum_chars = sum(1 for c in line if c.isalnum() or c.isspace() or c in ",.-_@()[]{}")
                if alnum_chars / len(line) > 0.85:
                    cleaned_line = re.sub(r"\s+", " ", line)
                    lines.append(cleaned_line)
                    
            # Deduplicate contiguous identical lines
            unique_lines = []
            for l in lines:
                if not unique_lines or unique_lines[-1] != l:
                    unique_lines.append(l)
                    
            return "\n".join(unique_lines)
        except Exception as ex:
            return f"Failed to parse DOC file via fallback: {ex}"

def parse_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".doc":
        return parse_doc(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
